"""Text extraction from various file formats for resume parsing."""

import io
import re
from typing import Tuple, Optional, Dict, Any
from pathlib import Path

import PyPDF2
from docx import Document
import structlog

from ..utils.logger import get_logger

logger = get_logger(__name__)

class TextExtractor:
    """Extracts text from PDF, DOCX, and TXT files with validation and cleaning."""
    
    # Supported file types and their MIME types
    SUPPORTED_TYPES = {
        'pdf': ['application/pdf'],
        'docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
        'txt': ['text/plain']
    }
    
    # Maximum file size (5MB)
    MAX_FILE_SIZE = 5 * 1024 * 1024
    
    # Minimum content requirement (100 words)
    MIN_WORD_COUNT = 100
    
    def __init__(self):
        """Initialize the TextExtractor."""
        self.logger = logger
    
    def extract(self, file_content: bytes, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from file content with metadata.
        
        Args:
            file_content: Raw file bytes
            file_type: File extension (pdf, docx, txt)
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            ValueError: For unsupported file types or validation failures
            RuntimeError: For extraction failures
        """
        self.logger.info("Starting text extraction", file_type=file_type, content_size=len(file_content))
        
        # Validate file type
        if file_type.lower() not in self.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Validate file size
        if len(file_content) > self.MAX_FILE_SIZE:
            raise ValueError(f"File size {len(file_content)} exceeds maximum {self.MAX_FILE_SIZE}")
        
        # Validate minimum content
        if len(file_content) == 0:
            raise ValueError("File is empty")
        
        metadata = {
            'file_type': file_type.lower(),
            'file_size': len(file_content),
            'extraction_method': None,
            'encoding': None,
            'word_count': 0,
            'extraction_errors': []
        }
        
        try:
            # Extract text based on file type
            if file_type.lower() == 'pdf':
                text = self._extract_pdf(file_content)
                metadata['extraction_method'] = 'pypdf2'
            elif file_type.lower() == 'docx':
                text = self._extract_docx(file_content)
                metadata['extraction_method'] = 'python-docx'
            elif file_type.lower() == 'txt':
                text, encoding = self._extract_txt(file_content)
                metadata['extraction_method'] = 'text-encoding-detection'
                metadata['encoding'] = encoding
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Clean and normalize text
            cleaned_text = self._clean_text(text)
            
            # Validate minimum word count
            word_count = len(cleaned_text.split())
            metadata['word_count'] = word_count
            
            if word_count < self.MIN_WORD_COUNT:
                raise ValueError(f"Extracted content has {word_count} words, minimum required is {self.MIN_WORD_COUNT}")
            
            self.logger.info("Text extraction completed", 
                           word_count=word_count, 
                           file_type=file_type)
            
            return cleaned_text, metadata
            
        except Exception as e:
            self.logger.error("Text extraction failed", 
                            error=str(e), 
                            file_type=file_type)
            metadata['extraction_errors'].append(str(e))
            raise RuntimeError(f"Text extraction failed: {str(e)}")
    
    def _extract_pdf(self, content: bytes) -> str:
        """
        Extract text from PDF content.
        
        Args:
            content: PDF file bytes
            
        Returns:
            Extracted text string
            
        Raises:
            RuntimeError: For PDF processing errors
        """
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                raise RuntimeError("Password-protected PDFs are not supported")
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    self.logger.warning("Failed to extract text from page", 
                                      page_num=page_num, 
                                      error=str(e))
                    continue
            
            if not text_parts:
                raise RuntimeError("No text could be extracted from PDF")
            
            return '\n'.join(text_parts)
            
        except PyPDF2.PdfReadError as e:
            raise RuntimeError(f"Invalid or corrupted PDF file: {str(e)}")
        except Exception as e:
            raise RuntimeError(f"PDF extraction failed: {str(e)}")
    
    def _extract_docx(self, content: bytes) -> str:
        """
        Extract text from DOCX content.
        
        Args:
            content: DOCX file bytes
            
        Returns:
            Extracted text string
            
        Raises:
            RuntimeError: For DOCX processing errors
        """
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            if not text_parts:
                raise RuntimeError("No text could be extracted from DOCX")
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            raise RuntimeError(f"DOCX extraction failed: {str(e)}")
    
    def _extract_txt(self, content: bytes) -> Tuple[str, str]:
        """
        Extract text from TXT content with encoding detection.
        
        Args:
            content: TXT file bytes
            
        Returns:
            Tuple of (extracted_text, detected_encoding)
            
        Raises:
            RuntimeError: For text processing errors
        """
        encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
        
        for encoding in encodings:
            try:
                text = content.decode(encoding)
                return text, encoding
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, try with error handling
        try:
            text = content.decode('utf-8', errors='replace')
            return text, 'utf-8-with-replace'
        except Exception as e:
            raise RuntimeError(f"Text file encoding detection failed: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned and normalized text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize unicode characters
        text = text.encode('ascii', 'ignore').decode('ascii')
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def validate_file_type(self, file_type: str) -> bool:
        """
        Validate if file type is supported.
        
        Args:
            file_type: File extension
            
        Returns:
            True if supported, False otherwise
        """
        return file_type.lower() in self.SUPPORTED_TYPES
    
    def get_supported_types(self) -> Dict[str, list]:
        """
        Get supported file types and their MIME types.
        
        Returns:
            Dictionary of supported file types
        """
        return self.SUPPORTED_TYPES.copy()
